"""
ResumeEnhancerService
---------------------
Two responsibilities:
  1. enhance()  → calls OpenAI to rewrite the resume with selected improvements
  2. to_docx()  → converts plain-text resume to a professional ATS-optimized DOCX
                  using a single enforced layout (best for ATS parsing)

Layout principles enforced on every download:
  - Single column, no tables, no text boxes, no graphics
  - Calibri font throughout (universally supported, ATS-safe)
  - Name: 18pt bold, centered
  - Contact line: 10pt, centered, separated by  |
  - Section headers: 11pt bold ALL CAPS with bottom border line
  - Body text: 10.5pt, left-aligned
  - Bullet points: real Word list bullets (not unicode characters)
  - Margins: 1 inch all sides
  - No colors, no icons, no columns

Env vars required:
  OPENAI_API_KEY  — from https://platform.openai.com/api-keys
"""

import io
import logging
import os
import re
from typing import List, Dict, Any

import httpx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()
logger = logging.getLogger(__name__)

OPENAI_MODEL    = "gpt-4o-mini"
OPENAI_BASE_URL = "https://api.openai.com/v1/chat/completions"

# Standard ATS section order — used to sort sections when building DOCX
SECTION_ORDER = [
    "CONTACT", "SUMMARY", "OBJECTIVE", "PROFILE",
    "EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EMPLOYMENT",
    "SKILLS", "TECHNICAL SKILLS", "CORE COMPETENCIES",
    "EDUCATION",
    "CERTIFICATIONS", "CERTIFICATES", "LICENSES",
    "PROJECTS",
    "AWARDS", "ACHIEVEMENTS", "HONORS",
    "PUBLICATIONS", "VOLUNTEER", "LANGUAGES",
]


class ResumeEnhancerService:

    def __init__(self):
        self.api_key = os.getenv("OPENAI_API_KEY", "")
        if not self.api_key:
            logger.warning("OPENAI_API_KEY is not set — enhancement calls will fail.")

    # ─────────────────────────────────────────────────────────────
    # 1.  AI ENHANCEMENT
    # ─────────────────────────────────────────────────────────────

    async def enhance(
        self,
        resume_text:     str,
        job_description: str,
        suggestions:     List[Dict[str, Any]],
    ) -> str:
        system = (
            "You are an elite professional resume writer and ATS optimization expert "
            "with 15+ years of experience helping candidates land interviews at top companies. "
            "Your task is to rewrite the provided resume, applying ALL the listed improvements "
            "while following these strict rules:\n\n"
            "PRESERVATION RULES (never violate these):\n"
            "- Keep ALL factual information exactly as-is: dates, company names, job titles, "
            "  universities, degrees, GPA, certifications, personal name, contact info\n"
            "- Never invent or fabricate any experience, skills, or achievements\n"
            "- Never remove any jobs, projects, or educational entries\n\n"
            "IMPROVEMENT RULES (apply all of these):\n"
            "- Replace weak verbs (managed, helped, worked on, assisted) with powerful action verbs "
            "  (engineered, spearheaded, optimized, architected, delivered, accelerated)\n"
            "- Add quantifiable metrics wherever plausible based on context "
            "  (e.g. 'managed team' → 'led cross-functional team of 6 engineers')\n"
            "- Weave in relevant keywords naturally throughout the text\n"
            "- Ensure each bullet point starts with a strong action verb\n"
            "- Tighten language — remove filler words and passive constructions\n"
            "- Ensure the summary/objective is compelling and keyword-rich\n\n"
            "OUTPUT FORMAT:\n"
            "Return ONLY the improved resume as clean plain text. Use this exact structure:\n"
            "- First line: candidate's full name (no label)\n"
            "- Second line: contact details separated by  |  (email | phone | location | linkedin)\n"
            "- Section headers in ALL CAPS on their own line (e.g. SUMMARY, EXPERIENCE, SKILLS)\n"
            "- Bullet points using the • character\n"
            "- Job entries: Company Name | Job Title | Start Date – End Date\n"
            "- Education entries: Institution | Degree | Graduation Year\n"
            "No preamble, no commentary, no markdown, no JSON — just the resume text."
        )

        has_jd   = bool(job_description.strip())
        jd_block = (
            f"\nTARGET JOB DESCRIPTION (optimize the resume for this role):\n"
            f"{'='*60}\n{job_description}\n{'='*60}\n"
            if has_jd else ""
        )

        # Format suggestions as detailed instructions
        improvements_text = "\n\n".join(
            f"IMPROVEMENT {i+1} [{s['category']} — {s['priority'].upper()} PRIORITY]\n"
            f"Problem: {s['issue']}\n"
            f"How to fix: {s['fix']}"
            + (f"\nExample: {s['example']}" if s.get("example") else "")
            for i, s in enumerate(suggestions)
        )

        user = (
            f"ORIGINAL RESUME:\n{'='*60}\n{resume_text}\n{'='*60}"
            f"{jd_block}\n\n"
            f"APPLY ALL OF THESE IMPROVEMENTS:\n{'='*60}\n{improvements_text}\n{'='*60}\n\n"
            f"Now rewrite the complete resume applying every improvement above. "
            f"The result should be significantly stronger than the original — "
            f"more impactful language, better keyword density, and clearer structure. "
            f"Remember: preserve all facts, improve only the presentation."
        )

        return await self._call_openai(system, user)

    # ─────────────────────────────────────────────────────────────
    # 2.  PROFESSIONAL ATS-OPTIMIZED DOCX GENERATION
    # ─────────────────────────────────────────────────────────────

    async def to_docx(self, resume_text: str) -> bytes:
        """
        Convert plain-text resume to a professionally formatted DOCX
        with a single enforced ATS-optimized layout.
        """
        doc = Document()
        self._set_margins(doc)
        self._set_default_font(doc)

        lines = [l.rstrip() for l in resume_text.split("\n")]
        lines = self._clean_lines(lines)

        i = 0
        name_done    = False
        contact_done = False

        while i < len(lines):
            line = lines[i].strip()

            # ── Empty line ─────────────────────────────────────
            if not line:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                i += 1
                continue

            # ── Candidate name (very first non-empty line) ─────
            if not name_done:
                self._add_name(doc, line)
                name_done = True
                i += 1
                continue

            # ── Contact info (second non-empty line) ──────────
            if not contact_done and self._is_contact_line(line):
                self._add_contact(doc, line)
                contact_done = True
                i += 1
                continue

            # ── Section header (ALL CAPS, short line) ─────────
            if self._is_section_header(line):
                self._add_section_header(doc, line)
                i += 1
                continue

            # ── Bullet point ───────────────────────────────────
            if self._is_bullet(line):
                self._add_bullet(doc, self._clean_bullet(line))
                i += 1
                continue

            # ── Job / Education entry line ─────────────────────
            if self._is_entry_line(line):
                self._add_entry_line(doc, line)
                i += 1
                continue

            # ── Regular paragraph ──────────────────────────────
            self._add_body_text(doc, line)
            i += 1

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ─────────────────────────────────────────────────────────────
    # DOCX HELPER — Document setup
    # ─────────────────────────────────────────────────────────────

    @staticmethod
    def _set_margins(doc: Document):
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

    @staticmethod
    def _set_default_font(doc: Document):
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # ─────────────────────────────────────────────────────────────
    # DOCX HELPER — Paragraph types
    # ─────────────────────────────────────────────────────────────

    @staticmethod
    def _add_name(doc: Document, text: str):
        """Candidate name — 18pt bold centered."""
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text.strip())
        run.bold      = True
        run.font.name = "Calibri"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    @staticmethod
    def _add_contact(doc: Document, text: str):
        """Contact line — 10pt centered, parts separated by  |  ."""
        # Normalize separators
        normalized = re.sub(r"\s*[|•·✦]\s*", "  |  ", text.strip())
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
        run = p.add_run(normalized)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    @staticmethod
    def _add_section_header(doc: Document, text: str):
        """
        Section header — 11pt bold ALL CAPS with a bottom border line.
        This is the most ATS-critical formatting element.
        """
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text.upper().strip())
        run.bold      = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

        # Add bottom border — thin single line, ATS-safe (it's just a paragraph border)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "4")       # half-points, 4 = 0.5pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")  # light grey
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _add_entry_line(doc: Document, text: str):
        """
        Job/Education entry header line.
        Bold company/institution, regular role and dates.
        Expected format: Company | Role | Dates  OR  Institution | Degree | Year
        """
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(1)

        parts = re.split(r"\s*\|\s*", text.strip(), maxsplit=2)
        if len(parts) >= 2:
            run1 = p.add_run(parts[0].strip())
            run1.bold      = True
            run1.font.name = "Calibri"
            run1.font.size = Pt(10.5)

            for part in parts[1:]:
                sep = p.add_run("  |  ")
                sep.font.name = "Calibri"
                sep.font.size = Pt(10.5)
                sep.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

                run = p.add_run(part.strip())
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)
        else:
            run = p.add_run(text.strip())
            run.bold      = True
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)

    @staticmethod
    def _add_bullet(doc: Document, text: str):
        """
        Bullet point using Word's built-in List Bullet style.
        Real Word bullets — not unicode • characters — for maximum ATS compatibility.
        """
        try:
            p = doc.add_paragraph(style="List Bullet")
        except KeyError:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)

        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)

    @staticmethod
    def _add_body_text(doc: Document, text: str):
        """Regular body paragraph."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)

    # ─────────────────────────────────────────────────────────────
    # DOCX HELPER — Line classifiers
    # ─────────────────────────────────────────────────────────────

    @staticmethod
    def _is_section_header(line: str) -> bool:
        """
        A line is a section header if it is:
        - ALL CAPS (ignoring spaces, colons, slashes, &, digits)
        - Reasonably short (under 45 chars)
        - Not a name or contact line
        """
        stripped = line.strip()
        if not stripped or len(stripped) > 45:
            return False
        alpha_only = re.sub(r"[^A-Za-z]", "", stripped)
        if not alpha_only:
            return False
        return alpha_only.isupper() and len(alpha_only) >= 3

    @staticmethod
    def _is_contact_line(line: str) -> bool:
        """Contact line typically contains email, phone, or linkedin."""
        low = line.lower()
        return (
            "@" in line
            or re.search(r"\+?\d[\d\s\-().]{7,}", line) is not None
            or "linkedin" in low
            or "github" in low
            or re.search(r"\||\s•\s", line) is not None
        )

    @staticmethod
    def _is_bullet(line: str) -> bool:
        """Detect bullet points."""
        stripped = line.strip()
        return (
            stripped.startswith(("•", "-", "–", "*", "·", "▪", "◦", "○"))
            or re.match(r"^\s{2,}[•\-\–\*]", line) is not None
        )

    @staticmethod
    def _is_entry_line(line: str) -> bool:
        """
        Job/education entry lines typically have | separators and date patterns.
        """
        has_pipe = "|" in line
        has_date = bool(re.search(
            r"(19|20)\d{2}|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Present|Current",
            line, re.IGNORECASE
        ))
        return has_pipe and has_date

    # ─────────────────────────────────────────────────────────────
    # DOCX HELPER — Text utilities
    # ─────────────────────────────────────────────────────────────

    @staticmethod
    def _clean_bullet(line: str) -> str:
        """Remove bullet character prefix — Word's List Bullet style adds its own."""
        return re.sub(r"^[\s•\-–\*·▪◦○]+", "", line).strip()

    @staticmethod
    def _clean_lines(lines: list) -> list:
        """Remove excessive blank lines (max 1 consecutive blank)."""
        cleaned    = []
        prev_blank = False
        for line in lines:
            is_blank = not line.strip()
            if is_blank and prev_blank:
                continue
            cleaned.append(line)
            prev_blank = is_blank
        return cleaned

    # ─────────────────────────────────────────────────────────────
    # OpenAI HTTP call
    # ─────────────────────────────────────────────────────────────

    async def _call_openai(self, system_instruction: str, user_prompt: str) -> str:
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type":  "application/json",
        }
        body = {
            "model":    OPENAI_MODEL,
            "messages": [
                {"role": "system", "content": system_instruction},
                {"role": "user",   "content": user_prompt},
            ],
            "temperature": 0.4,
            "max_tokens":  4096,
        }
        async with httpx.AsyncClient(timeout=120) as client:
            resp = await client.post(OPENAI_BASE_URL, headers=headers, json=body)
            if resp.status_code != 200:
                detail = resp.json().get("error", {}).get("message", resp.text)
                raise RuntimeError(f"OpenAI API {resp.status_code}: {detail}")
            data = resp.json()
        return data["choices"][0]["message"]["content"].strip()
